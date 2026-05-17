"""Fix code formatting in existing reports without overwriting other content."""
import sys
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

sys.stdout.reconfigure(encoding='utf-8')

REPORTS = [
    "C:/Users/50469/python-learning/实验七_Vue3待办事项应用_实验报告.docx",
    "C:/Users/50469/python-learning/实验八_Vue3响应式商城_实验报告.docx",
]

def has_code_formatting(paragraph):
    """Check if a paragraph already has the gray shading."""
    pPr = paragraph._element.find(qn('w:pPr'))
    if pPr is not None:
        shd = pPr.find(qn('w:shd'))
        if shd is not None and shd.get(qn('w:fill')) == 'F2F2F2':
            return True
    return False

def apply_code_formatting(paragraph):
    """Apply code block shading and border."""
    pPr = paragraph._element.get_or_add_pPr()

    # Shading
    shd_elem = pPr.find(qn('w:shd'))
    if shd_elem is None:
        shd_elem = OxmlElement('w:shd')
        pPr.append(shd_elem)
    shd_elem.set(qn('w:val'), 'clear')
    shd_elem.set(qn('w:fill'), 'F2F2F2')

    # Border
    existing_bdr = pPr.find(qn('w:pBdr'))
    if existing_bdr is None:
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)
    else:
        pBdr = existing_bdr
    for side in ('top', 'left', 'bottom', 'right'):
        # Check if border side already exists
        found = False
        for child in pBdr:
            if child.tag == qn(f'w:{side}'):
                found = True
                break
        if not found:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '4' if side in ('left', 'right') else '1')
            border.set(qn('w:color'), 'auto')
            pBdr.append(border)

def is_code_line(text):
    """Detect if a trimmed line looks like code."""
    stripped = text.strip()
    if not stripped:
        return False

    # HTML / template tags
    if stripped.startswith('<') and '>' in stripped:
        return True
    if stripped.startswith('</'):
        return True

    # JS/TS keywords and patterns
    js_starters = (
        'import ', 'export ', 'const ', 'let ', 'var ',
        'function ', 'async ', 'await ', 'return ',
        'if ', 'else ', 'else if ',
        'try ', 'catch ', 'finally ',
        'throw ', 'new ', 'typeof ', 'delete ',
        'switch ', 'case ', 'break ', 'continue ',
        'for ', 'while ', 'do ', 'in ',
        'class ', 'extends ', 'interface ',
        'default ', 'from ', 'yield ',
        'import', 'export',
    )
    if any(stripped.startswith(p) for p in js_starters):
        return True

    # Vue directives / template expressions
    if stripped.startswith('{{') or stripped.startswith('v-'):
        return True
    if ':key=' in stripped or ':src=' in stripped or ':class=' in stripped:
        return True
    if 'v-if' in stripped or 'v-else' in stripped or 'v-for' in stripped or 'v-model' in stripped:
        return True
    if 'v-bind' in stripped or 'v-on' in stripped or 'v-show' in stripped:
        return True

    # CSS / style patterns
    if stripped.startswith('@media'):
        return True
    if stripped.startswith('* {'):
        return True
    if stripped.startswith('.') and ('{' in stripped or stripped.endswith(',') or stripped.endswith(' {')):
        return True
    if text.lstrip().startswith('.') and ':' in text and text.rstrip().endswith(';'):
        return True
    if 'flex' in stripped and ';' in stripped:
        return True
    if stripped.endswith(';') and ':' in stripped and not any(c.isascii() and c.isalpha() for c in stripped):
        return True

    # JS object / array patterns
    if stripped.startswith('{') or stripped.startswith('}'):
        return True
    if stripped.startswith('[') or stripped.startswith(']'):
        return True
    if stripped.startswith('(') or stripped.startswith(')'):
        return True
    if stripped.startswith('...'):  # spread operator
        return True

    # Lines ending with specific code patterns
    if stripped in ('});', '})', ');', ']', '};', '},', '],', ')}', '))'):
        return True
    if stripped.endswith('),') or stripped.endswith(',') and not stripped.endswith('。'):
        return True

    # Property access / method call
    if '.value' in stripped or '.text' in stripped:
        return True
    if '.push(' in stripped or '.filter(' in stripped or '.map(' in stripped:
        return True
    if '.find(' in stripped or '.splice(' in stripped or '.findIndex(' in stripped:
        return True
    if '.length' in stripped or '.trim()' in stripped:
        return True
    if '.getItem(' in stripped or '.setItem(' in stripped:
        return True
    if 'document.' in stripped or 'window.' in stripped:
        return True
    if 'console.' in stripped:
        return True
    if 'Promise.' in stripped or 'setTimeout' in stripped:
        return True
    if ' => ' in stripped or '=>' in stripped:
        return True
    if '`${' in stripped:
        return True

    # Arrow function / function call
    if stripped.startswith('}') and stripped.endswith(')'):
        return True

    # JS object property (key: value pattern in code context)
    if ': ' in stripped and stripped.endswith(',') and not stripped.endswith('。'):
        parts = stripped.split(': ', 1)
        key_part = parts[0].strip()
        # Key is typically a simple identifier or string
        if key_part and (key_part.isidentifier() or key_part.startswith("'") or key_part.startswith('"')):
            return True
        if key_part.startswith('...'):
            return True

    # Pipe-style code patterns
    if stripped.count('|') >= 2:
        return True

    return False

for path in REPORTS:
    print(f"\nProcessing: {path}")
    doc = Document(path)
    cell = doc.tables[0].rows[3].cells[0]

    fixed = 0
    already = 0
    for para in cell.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        looks_like_code = is_code_line(text)
        already_formatted = has_code_formatting(para)

        if looks_like_code and not already_formatted:
            apply_code_formatting(para)
            fixed += 1
        elif looks_like_code and already_formatted:
            already += 1

    doc.save(path)
    print(f"  Fixed: {fixed} code lines (already formatted: {already})")

print("\nDone! All existing reports updated.")
